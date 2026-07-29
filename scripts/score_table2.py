#!/usr/bin/env python3
"""
Table 2 evaluation for RiskScribe layoutbank pairs.

Scores infographic_decoration.png vs original_preview.png on:
  - Layout Validity (VLM checklist on decoration)
  - Element coverage vs original (VLM on both)
  - Aesthetics win/tie/loss (blinded pairwise VLM)
  - Readability win/tie/loss (blinded pairwise VLM)
  - Referenced scores: (100*N_win + 50*N_tie) / N

Outputs:
  - results/table2/per_case_results.jsonl  (resume-friendly)
  - results/table2/per_case_results.csv
  - results/table2/aggregate_summary.json
  - updates RiskScribe_Final_Score_Tables.docx Table 2 (RiskScribe column)
"""

from __future__ import annotations

import base64
import csv
import io
import json
import os
import random
import re
import sys
import time
from pathlib import Path

from openai import OpenAI
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
LAYOUTBANK = ROOT / "data" / "table2"
OUT_DIR = ROOT / "results" / "table2"
JSONL_PATH = OUT_DIR / "per_case_results.jsonl"
CSV_PATH = OUT_DIR / "per_case_results.csv"
SUMMARY_PATH = OUT_DIR / "aggregate_summary.json"
DOCX_PATH = ROOT / "results" / "RiskScribe_Final_Score_Tables.docx"
API_KEYS_PATH = ROOT / "api_keys.txt"

MODEL = os.environ.get("TABLE2_VLM_MODEL", "gpt-4o")
MAX_IMAGE_SIDE = 1600  # downscale long edge for cost/latency
MAX_RETRIES = 5
SEED = 42

# ---------------------------------------------------------------------------
# API key
# ---------------------------------------------------------------------------

def load_openai_key() -> str:
    env = os.environ.get("OPENAI_API_KEY")
    if env:
        return env.strip()
    text = API_KEYS_PATH.read_text()
    # Prefer explicit GPT: line, else first sk- key
    for line in text.splitlines():
        s = line.strip()
        if s.lower().startswith("gpt:") or s.lower().startswith("openai:"):
            # key may be on same or next non-empty line
            parts = s.split(":", 1)
            if len(parts) == 2 and parts[1].strip().startswith("sk-"):
                return parts[1].strip()
        if s.startswith("sk-"):
            return s
    # multi-line: "GPT:" then key on next line
    lines = [ln.strip() for ln in text.splitlines()]
    for i, ln in enumerate(lines):
        if ln.lower().rstrip(":") in {"gpt", "openai"} and i + 1 < len(lines):
            if lines[i + 1].startswith("sk-"):
                return lines[i + 1]
    raise RuntimeError("Could not find OpenAI API key in env or api_keys.txt")


# ---------------------------------------------------------------------------
# Image helpers
# ---------------------------------------------------------------------------

def encode_image(path: Path) -> tuple[str, str]:
    """Return (media_type, base64_data) after optional downscale."""
    img = Image.open(path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    w, h = img.size
    scale = min(1.0, MAX_IMAGE_SIDE / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85, optimize=True)
    b64 = base64.standard_b64encode(buf.getvalue()).decode("ascii")
    return "image/jpeg", b64


def image_content(path: Path, detail: str = "high") -> dict:
    media, b64 = encode_image(path)
    return {
        "type": "image_url",
        "image_url": {
            "url": f"data:{media};base64,{b64}",
            "detail": detail,
        },
    }


# ---------------------------------------------------------------------------
# VLM call
# ---------------------------------------------------------------------------

def extract_json(text: str) -> dict:
    text = text.strip()
    # strip markdown fences if present
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if fence:
        text = fence.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # try first {...} block
        m = re.search(r"\{[\s\S]*\}", text)
        if not m:
            raise
        return json.loads(m.group(0))


def vlm_json(client: OpenAI, system: str, user_parts: list, max_tokens: int = 1200) -> dict:
    last_err = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                temperature=0,
                max_tokens=max_tokens,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_parts},
                ],
            )
            content = resp.choices[0].message.content or "{}"
            return extract_json(content)
        except Exception as e:
            last_err = e
            wait = min(2 ** attempt, 30) + random.uniform(0, 1)
            print(f"    retry {attempt}/{MAX_RETRIES} after error: {e} (sleep {wait:.1f}s)", flush=True)
            time.sleep(wait)
    raise RuntimeError(f"VLM call failed after {MAX_RETRIES} retries: {last_err}")


# ---------------------------------------------------------------------------
# Metric prompts
# ---------------------------------------------------------------------------

LAYOUT_SYSTEM = """You are evaluating the LAYOUT VALIDITY of a public-safety infographic image.
Judge visual/layout structure only — NOT factual correctness of numbers or claims.

Checklist (each true/false):
1. full_render: entire canvas is visible and complete (not cut off mid-design)
2. no_text_overflow: text stays inside its containers / does not run off edges
3. no_illegal_overlap: text/icons do not illegibly stack on each other
4. no_clipping: important content is not cropped at edges
5. structure_legible: hierarchy is readable (title vs body vs actions distinguishable)

A case PASSES only if ALL five are true.

Return ONLY JSON:
{
  "full_render": true/false,
  "no_text_overflow": true/false,
  "no_illegal_overlap": true/false,
  "no_clipping": true/false,
  "structure_legible": true/false,
  "pass": true/false,
  "one_line_reason": "short reason"
}"""


def score_layout(client: OpenAI, decoration: Path) -> dict:
    parts = [
        {"type": "text", "text": "Score layout validity for this RiskScribe-generated public-safety infographic."},
        image_content(decoration),
    ]
    return vlm_json(client, LAYOUT_SYSTEM, parts, max_tokens=500)


COVERAGE_SYSTEM = """You compare a human ORIGINAL public-agency infographic (Image A) to a GENERATED remake (Image B).

Task:
1. List the ORIGINAL's distinct story/content elements (Image A only).
   Elements are communicative units such as: main title/headline, subtitle, numbered steps,
   bullet tips, key statistic/KPI, warning callout, map/diagram region, icon-labeled action,
   source/footer/URL, logo/agency mark, checklist items, definitions (e.g. watch vs warning).
   Merge tiny fragments; aim for 4-12 substantial elements.
2. For each element, set present_in_generated true if Image B conveys the SAME communicative
   content (paraphrase OK). Do not require identical layout, color, or wording.
3. coverage = number present_in_generated / number of original elements (0-1).

Return ONLY JSON:
{
  "elements": [
    {"id": 1, "description": "...", "present_in_generated": true/false}
  ],
  "n_original": <int>,
  "n_matched": <int>,
  "coverage": <float 0-1>,
  "one_line_reason": "short"
}"""


def score_coverage(client: OpenAI, original: Path, decoration: Path) -> dict:
    parts = [
        {
            "type": "text",
            "text": (
                "Image A = ORIGINAL agency infographic.\n"
                "Image B = GENERATED remake (RiskScribe decoration).\n"
                "Compute element coverage of B vs A."
            ),
        },
        {"type": "text", "text": "IMAGE A (original):"},
        image_content(original),
        {"type": "text", "text": "IMAGE B (generated):"},
        image_content(decoration),
    ]
    data = vlm_json(client, COVERAGE_SYSTEM, parts, max_tokens=1500)
    # recompute coverage from elements if present
    elements = data.get("elements") or []
    if elements:
        n_orig = len(elements)
        n_match = sum(1 for e in elements if e.get("present_in_generated"))
        data["n_original"] = n_orig
        data["n_matched"] = n_match
        data["coverage"] = round(n_match / n_orig, 4) if n_orig else 0.0
    else:
        data["coverage"] = float(data.get("coverage") or 0.0)
        data["n_original"] = int(data.get("n_original") or 0)
        data["n_matched"] = int(data.get("n_matched") or 0)
    return data


PAIRWISE_AESTHETICS_SYSTEM = """You are rating visual AESTHETIC QUALITY of two public-safety infographics.
Do NOT verify factual correctness. Judge only: visual appeal, hierarchy, balance, use of space,
professionalism, color/composition quality.

You will see Image A and Image B (order is random; do not assume which is original).

Choose:
- "A" if Image A is clearly better aesthetically
- "B" if Image B is clearly better aesthetically
- "tie" if roughly equal (minor differences only)

Return ONLY JSON:
{"winner": "A"|"B"|"tie", "one_line_reason": "short"}"""


PAIRWISE_READABILITY_SYSTEM = """You are rating READABILITY of two public-safety infographics.
Do NOT verify factual correctness. Judge only: text legibility, contrast, label clarity,
whether a reader can quickly scan the message under realistic viewing conditions.

You will see Image A and Image B (order is random).

Choose:
- "A" if Image A is clearly more readable
- "B" if Image B is clearly more readable
- "tie" if roughly equal

Return ONLY JSON:
{"winner": "A"|"B"|"tie", "one_line_reason": "short"}"""


def score_pairwise(
    client: OpenAI,
    system: str,
    original: Path,
    decoration: Path,
    rng: random.Random,
) -> dict:
    """Blinded pairwise: randomize A/B, map winner back to gen_win|tie|gen_loss."""
    gen_is_a = rng.random() < 0.5
    if gen_is_a:
        path_a, path_b = decoration, original
        label_a, label_b = "generated", "original"
    else:
        path_a, path_b = original, decoration
        label_a, label_b = "original", "generated"

    parts = [
        {"type": "text", "text": "Compare Image A and Image B. Return JSON with winner A, B, or tie."},
        {"type": "text", "text": "IMAGE A:"},
        image_content(path_a),
        {"type": "text", "text": "IMAGE B:"},
        image_content(path_b),
    ]
    raw = vlm_json(client, system, parts, max_tokens=300)
    winner = str(raw.get("winner", "tie")).strip().upper()
    if winner not in {"A", "B", "TIE"}:
        # normalize
        w = winner.lower()
        if "tie" in w or w in {"equal", "same"}:
            winner = "TIE"
        elif w.startswith("a"):
            winner = "A"
        elif w.startswith("b"):
            winner = "B"
        else:
            winner = "TIE"

    if winner == "TIE":
        verdict = "tie"
    elif winner == "A":
        verdict = "gen_win" if gen_is_a else "gen_loss"
    else:  # B
        verdict = "gen_win" if not gen_is_a else "gen_loss"

    return {
        "verdict": verdict,
        "raw_winner": winner,
        "gen_was_image": "A" if gen_is_a else "B",
        "label_a": label_a,
        "label_b": label_b,
        "one_line_reason": raw.get("one_line_reason", ""),
    }


# ---------------------------------------------------------------------------
# Case discovery / resume
# ---------------------------------------------------------------------------

def list_cases() -> list[Path]:
    cases = sorted(
        p for p in LAYOUTBANK.iterdir()
        if p.is_dir() and (p / "infographic_decoration.png").exists()
        and (p / "original_preview.png").exists()
    )
    return cases


def load_done_ids() -> set[str]:
    done = set()
    if not JSONL_PATH.exists():
        return done
    for line in JSONL_PATH.read_text().splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            rec = json.loads(line)
            if rec.get("case_id") and rec.get("ok"):
                done.add(rec["case_id"])
        except json.JSONDecodeError:
            continue
    return done


def append_jsonl(rec: dict) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    with JSONL_PATH.open("a") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")


# ---------------------------------------------------------------------------
# Aggregate + write outputs
# ---------------------------------------------------------------------------

def load_all_records() -> list[dict]:
    if not JSONL_PATH.exists():
        return []
    # last record per case_id wins
    by_id = {}
    for line in JSONL_PATH.read_text().splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            rec = json.loads(line)
        except json.JSONDecodeError:
            continue
        if rec.get("case_id") and rec.get("ok"):
            by_id[rec["case_id"]] = rec
    return [by_id[k] for k in sorted(by_id)]


def aggregate(records: list[dict]) -> dict:
    n = len(records)
    if n == 0:
        raise RuntimeError("No successful records to aggregate")

    layout_pass = sum(1 for r in records if r.get("layout_pass"))
    coverages = [float(r["element_coverage"]) for r in records]

    def counts(key: str):
        wins = sum(1 for r in records if r.get(key) == "gen_win")
        ties = sum(1 for r in records if r.get(key) == "tie")
        losses = sum(1 for r in records if r.get(key) == "gen_loss")
        return wins, ties, losses

    a_w, a_t, a_l = counts("aesthetics_verdict")
    r_w, r_t, r_l = counts("readability_verdict")

    def ref_score(w, t, l, n_total):
        return round((100 * w + 50 * t + 0 * l) / n_total, 1)

    summary = {
        "n_cases": n,
        "model": MODEL,
        "layout_validity_pass_rate_pct": round(100.0 * layout_pass / n, 1),
        "layout_pass_count": layout_pass,
        "element_coverage_mean": round(sum(coverages) / n, 3),
        "aesthetics": {
            "win": a_w,
            "tie": a_t,
            "loss": a_l,
            "win_pct": round(100.0 * a_w / n, 1),
            "tie_pct": round(100.0 * a_t / n, 1),
            "loss_pct": round(100.0 * a_l / n, 1),
            "referenced_score": ref_score(a_w, a_t, a_l, n),
        },
        "readability": {
            "win": r_w,
            "tie": r_t,
            "loss": r_l,
            "win_pct": round(100.0 * r_w / n, 1),
            "tie_pct": round(100.0 * r_t / n, 1),
            "loss_pct": round(100.0 * r_l / n, 1),
            "referenced_score": ref_score(r_w, r_t, r_l, n),
        },
        "formula": "referenced = (100*N_win + 50*N_tie + 0*N_loss) / N",
        "notes": (
            f"N={n} (available layoutbank pairs; table template planned for 80). "
            "GPT-Img-2 omitted (optional). Original anchor: coverage=1.00, "
            "referenced aesthetics/readability=50.0 by definition."
        ),
    }
    return summary


def write_csv(records: list[dict]) -> None:
    fields = [
        "case_id",
        "layout_pass",
        "element_coverage",
        "n_original_elements",
        "n_matched_elements",
        "aesthetics_verdict",
        "readability_verdict",
        "layout_reason",
        "coverage_reason",
        "aesthetics_reason",
        "readability_reason",
    ]
    with CSV_PATH.open("w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        for r in records:
            w.writerow({
                "case_id": r["case_id"],
                "layout_pass": int(bool(r.get("layout_pass"))),
                "element_coverage": r.get("element_coverage"),
                "n_original_elements": r.get("n_original_elements"),
                "n_matched_elements": r.get("n_matched_elements"),
                "aesthetics_verdict": r.get("aesthetics_verdict"),
                "readability_verdict": r.get("readability_verdict"),
                "layout_reason": r.get("layout_reason", ""),
                "coverage_reason": r.get("coverage_reason", ""),
                "aesthetics_reason": r.get("aesthetics_reason", ""),
                "readability_reason": r.get("readability_reason", ""),
            })


def fill_docx(summary: dict) -> None:
    """Fill RiskScribe column cells in Table 2 of the score tables docx."""
    from docx import Document

    if not DOCX_PATH.exists():
        print(f"WARNING: {DOCX_PATH} not found; skip docx fill", flush=True)
        return

    doc = Document(str(DOCX_PATH))
    # Table 2 is the second table in the document (index 1)
    if len(doc.tables) < 2:
        print(f"WARNING: expected >=2 tables, found {len(doc.tables)}", flush=True)
        return

    t = doc.tables[1]
    a = summary["aesthetics"]
    r = summary["readability"]

    # Row structure from earlier extraction (header row 0):
    # 0 Metric | Scale | Judge | RiskScribe | GPT-Img-2 | Original
    # 1 Layout Validity
    # 2 Element coverage
    # 3 Aesthetics win/tie/loss
    # 4 Aesthetics referenced score
    # 5 Readability referenced score
    values = {
        1: f"{summary['layout_validity_pass_rate_pct']}%",  # Layout Validity
        2: f"{summary['element_coverage_mean']:.2f}",  # Element coverage
        3: f"{a['win_pct']}% / {a['tie_pct']}% / {a['loss_pct']}%",  # win/tie/loss
        4: f"{a['referenced_score']}",  # aesthetics referenced
        5: f"{r['referenced_score']}",  # readability referenced
    }
    # Original column anchors
    original_vals = {
        1: "report for context",
        2: "1.00 by definition",
        3: "—",
        4: "50.0 by definition",
        5: "50.0 by definition",
    }

    for row_idx, val in values.items():
        if row_idx >= len(t.rows):
            continue
        row = t.rows[row_idx]
        # RiskScribe is column 3 (0-indexed) if 6 cols
        if len(row.cells) >= 4:
            row.cells[3].text = val
        if len(row.cells) >= 6 and row_idx in original_vals:
            # only set if empty-ish
            if not row.cells[5].text.strip() or "definition" in original_vals[row_idx].lower() or original_vals[row_idx] == "—":
                row.cells[5].text = original_vals[row_idx]
        # leave GPT-Img-2 (col 4) blank / dash
        if len(row.cells) >= 5:
            if not row.cells[4].text.strip():
                row.cells[4].text = "—"

    out_docx = OUT_DIR / "RiskScribe_Final_Score_Tables_filled.docx"
    doc.save(str(out_docx))
    # also overwrite main file with backup first
    backup = ROOT / "archive" / "RiskScribe_Final_Score_Tables.backup.docx"
    if not backup.exists():
        backup.write_bytes(DOCX_PATH.read_bytes())
    doc.save(str(DOCX_PATH))
    print(f"Filled Table 2 in {DOCX_PATH} (backup: {backup.name})", flush=True)
    print(f"Copy also at {out_docx}", flush=True)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def score_one(client: OpenAI, case_dir: Path, rng: random.Random) -> dict:
    case_id = case_dir.name
    decoration = case_dir / "infographic_decoration.png"
    original = case_dir / "original_preview.png"

    print(f"  [1/4] layout validity...", flush=True)
    layout = score_layout(client, decoration)
    layout_pass = bool(layout.get("pass"))
    # enforce pass = all checklist true
    checks = ["full_render", "no_text_overflow", "no_illegal_overlap", "no_clipping", "structure_legible"]
    if all(k in layout for k in checks):
        layout_pass = all(bool(layout[k]) for k in checks)

    print(f"  [2/4] element coverage...", flush=True)
    coverage = score_coverage(client, original, decoration)

    print(f"  [3/4] aesthetics pairwise...", flush=True)
    aesthetics = score_pairwise(client, PAIRWISE_AESTHETICS_SYSTEM, original, decoration, rng)

    print(f"  [4/4] readability pairwise...", flush=True)
    readability = score_pairwise(client, PAIRWISE_READABILITY_SYSTEM, original, decoration, rng)

    return {
        "case_id": case_id,
        "ok": True,
        "model": MODEL,
        "layout_pass": layout_pass,
        "layout_detail": layout,
        "layout_reason": layout.get("one_line_reason", ""),
        "element_coverage": coverage.get("coverage", 0.0),
        "n_original_elements": coverage.get("n_original", 0),
        "n_matched_elements": coverage.get("n_matched", 0),
        "coverage_detail": coverage,
        "coverage_reason": coverage.get("one_line_reason", ""),
        "aesthetics_verdict": aesthetics["verdict"],
        "aesthetics_detail": aesthetics,
        "aesthetics_reason": aesthetics.get("one_line_reason", ""),
        "readability_verdict": readability["verdict"],
        "readability_detail": readability,
        "readability_reason": readability.get("one_line_reason", ""),
    }


def main() -> int:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    key = load_openai_key()
    client = OpenAI(api_key=key)
    cases = list_cases()
    print(f"Found {len(cases)} layoutbank pairs", flush=True)
    print(f"Model: {MODEL}", flush=True)

    done = load_done_ids()
    print(f"Already completed: {len(done)}", flush=True)

    rng = random.Random(SEED)
    # pre-draw seeds per case for reproducible blinding even if resumed
    case_seeds = {c.name: rng.randint(0, 10**9) for c in cases}

    failures = []
    for i, case_dir in enumerate(cases, 1):
        case_id = case_dir.name
        if case_id in done:
            print(f"[{i}/{len(cases)}] SKIP {case_id} (done)", flush=True)
            continue
        print(f"[{i}/{len(cases)}] SCORE {case_id}", flush=True)
        case_rng = random.Random(case_seeds[case_id])
        try:
            rec = score_one(client, case_dir, case_rng)
            append_jsonl(rec)
            done.add(case_id)
            print(
                f"    -> layout_pass={rec['layout_pass']} "
                f"coverage={rec['element_coverage']} "
                f"aes={rec['aesthetics_verdict']} "
                f"read={rec['readability_verdict']}",
                flush=True,
            )
        except Exception as e:
            print(f"    FAILED: {e}", flush=True)
            append_jsonl({"case_id": case_id, "ok": False, "error": str(e)})
            failures.append((case_id, str(e)))
            # brief pause then continue
            time.sleep(2)

    records = load_all_records()
    print(f"\nSuccessful records: {len(records)} / {len(cases)}", flush=True)
    if len(records) == 0:
        print("No results; abort aggregate", flush=True)
        return 1

    summary = aggregate(records)
    SUMMARY_PATH.write_text(json.dumps(summary, indent=2))
    write_csv(records)
    fill_docx(summary)

    print("\n========== TABLE 2 AGGREGATE (RiskScribe) ==========", flush=True)
    print(json.dumps(summary, indent=2), flush=True)
    if failures:
        print(f"\nFailures ({len(failures)}):", flush=True)
        for cid, err in failures:
            print(f"  {cid}: {err}", flush=True)
    print("\nDone.", flush=True)
    return 0 if not failures else 2


if __name__ == "__main__":
    sys.exit(main())
