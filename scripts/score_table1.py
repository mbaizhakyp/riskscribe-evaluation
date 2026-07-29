#!/usr/bin/env python3
"""
Table 1 scoring for supplemental complex-data cases.

Systems scored (N=5 cases each):
  gpt_img_1_5, gpt_img_2, ci_sol, ci_luna, riskscribe

Metrics (Expert appropriateness SKIPPED; framework distance / ablation blank):
  C1 FFR, HR, Numeric Accuracy
  C2 Chart-rule pass (proxy heuristic)
  C3 Layout Validity, Story Completeness
  C4 Aesthetic Quality (VLM 1-5, median of 3)

Judge model: gpt-4o
RiskScribe scored on infographic_decoration.png only.
Baselines scored on generated PNGs.
Fidelity: VLM transcribes visible numbers -> deterministic compare to registry gold.
"""

from __future__ import annotations

import base64
import csv
import io
import json
import os
import random
import re
import statistics
import sys
import time
from pathlib import Path
from typing import Any

from openai import OpenAI
from PIL import Image

ROOT = Path(__file__).resolve().parent.parent
SUPP = ROOT / "data" / "table1"
GEN = ROOT / "results" / "table1" / "generations"
OUT = ROOT / "results" / "table1"
JSONL = OUT / "per_case_system_results.jsonl"
CSV_PATH = OUT / "per_case_system_results.csv"
SUMMARY = OUT / "aggregate_summary.json"
DOCX_PATH = ROOT / "results" / "RiskScribe_Final_Score_Tables.docx"
API_KEYS = ROOT / "api_keys.txt"

MODEL = os.environ.get("TABLE1_VLM_MODEL", "gpt-4o")
MAX_SIDE = 1600
MAX_RETRIES = 5
AESTHETIC_PASSES = 3
def discover_table1_cases() -> list[str]:
    """Discover case IDs under data/table1 that have required scoring artifacts."""
    if not SUPP.exists():
        return []
    cases = []
    for d in sorted(SUPP.iterdir()):
        if not d.is_dir() or d.name.startswith("_") or d.name.startswith("."):
            continue
        # require registry for gold + at least riskscribe image path may be generated later
        if (d / "immutable_registry.json").exists():
            cases.append(d.name)
    # optional override: config/table1_cases.json
    cfg = ROOT / "config" / "table1_cases.json"
    if cfg.exists():
        try:
            wanted = json.loads(cfg.read_text())
            if isinstance(wanted, list) and wanted:
                return [c for c in wanted if c in cases or (SUPP / c).exists()]
        except Exception:
            pass
    return cases


CASES = discover_table1_cases()
SYSTEMS = ["gpt_img_1_5", "gpt_img_2", "ci_sol", "ci_luna", "riskscribe"]

# Table column order matching score table header
SYSTEM_COL = {
    "gpt_img_1_5": "GPT-Img-1.5",
    "gpt_img_2": "GPT-Img-2",
    "ci_sol": "CI sol",
    "ci_luna": "CI luna",
    "riskscribe": "RiskScribe",
}


def load_key() -> str:
    env = os.environ.get("OPENAI_API_KEY")
    if env:
        return env.strip()
    for line in API_KEYS.read_text().splitlines():
        s = line.strip()
        if s.startswith("sk-"):
            return s
    raise RuntimeError("OpenAI key not found")


def image_path(case_id: str, system: str) -> Path:
    if system == "riskscribe":
        return SUPP / case_id / "infographic_decoration.png"
    return GEN / case_id / system / f"{case_id}_{system}.png"


def encode_image(path: Path) -> dict:
    img = Image.open(path)
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    w, h = img.size
    scale = min(1.0, MAX_SIDE / max(w, h))
    if scale < 1.0:
        img = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85, optimize=True)
    b64 = base64.standard_b64encode(buf.getvalue()).decode("ascii")
    return {
        "type": "image_url",
        "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "high"},
    }


def extract_json(text: str) -> dict:
    text = (text or "").strip()
    fence = re.search(r"```(?:json)?\s*([\s\S]*?)```", text)
    if fence:
        text = fence.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        m = re.search(r"\{[\s\S]*\}", text)
        if not m:
            raise
        return json.loads(m.group(0))


def vlm_json(client: OpenAI, system: str, user_parts: list, max_tokens: int = 2000) -> dict:
    last = None
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = client.chat.completions.create(
                model=MODEL,
                temperature=0,
                max_tokens=max_tokens,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system},
                    {"role": "user", "content": user_parts},
                ],
            )
            return extract_json(resp.choices[0].message.content or "{}")
        except Exception as e:
            last = e
            time.sleep(min(2 ** attempt, 30) + random.random())
    raise RuntimeError(f"VLM failed: {last}")


# ---------------------------------------------------------------------------
# Gold values from registry
# ---------------------------------------------------------------------------

def load_registry(case_id: str) -> dict:
    return json.loads((SUPP / case_id / "immutable_registry.json").read_text())


def gold_numeric_values(reg: dict) -> list[dict]:
    """Authoritative numeric values for fidelity matching."""
    gold = []
    for tid, target in (reg.get("data_targets") or {}).items():
        for i, row in enumerate(target.get("records") or []):
            for k, v in row.items():
                if isinstance(v, bool):
                    continue
                if isinstance(v, (int, float)):
                    gold.append(
                        {
                            "source": f"{tid}.{k}",
                            "row": i,
                            "value": float(v),
                            "field": k,
                            "target": tid,
                        }
                    )
    for f in (reg.get("fact_packet") or {}).get("facts") or []:
        v = f.get("value")
        if isinstance(v, (int, float)) and not isinstance(v, bool):
            gold.append(
                {
                    "source": f.get("fact_id"),
                    "row": None,
                    "value": float(v),
                    "field": f.get("fact_id"),
                    "target": "fact_packet",
                    "unit": f.get("unit"),
                    "display": f.get("display"),
                }
            )
        for tok in f.get("numeric_tokens") or []:
            try:
                gold.append(
                    {
                        "source": f"{f.get('fact_id')}.token",
                        "value": float(tok),
                        "field": "token",
                        "target": "fact_packet",
                    }
                )
            except ValueError:
                pass
    # de-dupe by rounded value+source loosely keep all for matching
    return gold


def values_match(shown: float, gold: float) -> bool:
    """Tolerance for OCR / display rounding."""
    if gold == 0:
        return abs(shown) < 1e-6
    # exact-ish for small numbers, relative for large
    if abs(gold) < 100:
        return abs(shown - gold) <= 0.05 + 1e-9  # 0.05 absolute for percents etc.
    return abs(shown - gold) / abs(gold) <= 0.005  # 0.5%


def score_fidelity(visible_values: list[dict], gold: list[dict]) -> dict:
    """
    FFR = fraction of visible numeric data values that match some gold value.
    HR = 1 - FFR
    Numeric accuracy = same as FFR here (exact-with-tolerance match rate of shown values).
    Illegible/missing extraction => empty visible => FFR treated as 0 if nothing extracted
    (protocol: illegible counts against FFR).
    """
    gold_vals = [g["value"] for g in gold]
    parsed = []
    for item in visible_values or []:
        raw = item.get("value")
        try:
            if isinstance(raw, str):
                raw = raw.replace(",", "").replace("%", "").strip()
            v = float(raw)
            parsed.append({"value": v, "unit": item.get("unit"), "label": item.get("label")})
        except (TypeError, ValueError):
            continue

    if not parsed:
        return {
            "ffr": 0.0,
            "hr": 1.0,
            "numeric_accuracy": 0.0,
            "n_visible": 0,
            "n_matched": 0,
            "n_hallucinated": 0,
            "matched_values": [],
            "hallucinated_values": [],
        }

    matched = []
    hallu = []
    for p in parsed:
        ok = any(values_match(p["value"], g) for g in gold_vals)
        if ok:
            matched.append(p)
        else:
            hallu.append(p)

    n = len(parsed)
    n_m = len(matched)
    ffr = n_m / n
    return {
        "ffr": round(ffr, 4),
        "hr": round(1.0 - ffr, 4),
        "numeric_accuracy": round(ffr, 4),
        "n_visible": n,
        "n_matched": n_m,
        "n_hallucinated": len(hallu),
        "matched_values": matched,
        "hallucinated_values": hallu,
    }


# ---------------------------------------------------------------------------
# Chart-rule proxy
# ---------------------------------------------------------------------------

def expected_chart_families(reg: dict) -> set[str]:
    """
    Proxy data-shape -> allowed chart families.
    Not the official Phase-2 engine; documented heuristic.
    """
    targets = set((reg.get("data_targets") or {}).keys())
    families: set[str] = set()
    if "land_cover_distribution" in targets:
        families |= {"treemap", "pie", "donut", "bar", "stacked_bar", "mosaic"}
    if "precipitation_history" in targets or "flood_history" in targets:
        families |= {"line", "area", "bar", "multi_line", "time_series"}
    if "river_gauge_history" in targets:
        families |= {"line", "scatter", "bubble", "bar", "dual_axis"}
    if "river_gauge_status" in targets and len(targets) == 1:
        # single status: KPI-only OK, or simple indicator
        families |= {"kpi_only", "none", "bar", "gauge", "indicator"}
    if "social_vulnerability_historical" in targets:
        families |= {"bar", "line", "table", "heatmap", "multi_panel"}
    if not families:
        families = {"bar", "line", "table", "kpi_only", "other"}
    return families


def normalize_chart_label(s: str) -> str:
    s = (s or "").lower().strip()
    s = s.replace("-", " ").replace("_", " ")
    mapping = [
        ("tree", "treemap"),
        ("pie", "pie"),
        ("donut", "donut"),
        ("line", "line"),
        ("area", "area"),
        ("scatter", "scatter"),
        ("bubble", "bubble"),
        ("stack", "stacked_bar"),
        ("bar", "bar"),
        ("column", "bar"),
        ("heatmap", "heatmap"),
        ("gauge", "gauge"),
        ("kpi", "kpi_only"),
        ("none", "none"),
        ("no chart", "none"),
        ("table", "table"),
        ("map", "map"),
        ("panel", "multi_panel"),
    ]
    for key, fam in mapping:
        if key in s:
            return fam
    return "other"


def chart_rule_pass(reg: dict, detected_charts: list[str]) -> int:
    allowed = expected_chart_families(reg)
    if not detected_charts:
        # KPI-only allowed only if status-like
        return 1 if "kpi_only" in allowed or "none" in allowed else 0
    for ch in detected_charts:
        fam = normalize_chart_label(ch)
        if fam in allowed or fam == "other":
            # "other" alone does not pass unless something allowed matched
            if fam in allowed:
                return 1
    # if any detected maps to allowed
    norms = {normalize_chart_label(c) for c in detected_charts}
    if norms & allowed:
        return 1
    # multi-target complex: multi_panel / map soft pass if map/panel present
    if "map" in norms and len(reg.get("data_targets") or {}) > 1:
        return 1
    return 0


# ---------------------------------------------------------------------------
# VLM prompts
# ---------------------------------------------------------------------------

EXTRACT_SYSTEM = """You extract structured content from a public-risk / data infographic image.
Do NOT judge whether numbers are correct — only transcribe what is visibly shown.

Return ONLY JSON:
{
  "visible_numbers": [
    {"value": <number without commas>, "unit": <string or null>, "label": <short context>}
  ],
  "charts_detected": ["treemap"|"pie"|"bar"|"line"|"scatter"|"bubble"|"area"|"table"|"map"|"kpi_only"|"none"|"other", ...],
  "layout": {
    "full_render": true/false,
    "no_text_overflow": true/false,
    "no_illegal_overlap": true/false,
    "no_clipping": true/false,
    "structure_legible": true/false,
    "pass": true/false,
    "one_line_reason": "short"
  },
  "story_elements": {
    "location": true/false,
    "time_or_period": true/false,
    "quantitative_evidence": true/false,
    "guidance_or_actions": true/false,
    "source": true/false
  },
  "story_completeness_count": <0-5 integer how many of the five are true>,
  "one_line_summary": "short"
}

Rules for visible_numbers:
- Include percentages, totals, axis tick values that encode data, KPI callouts, table cells.
- Prefer unique data values (skip pure decorative years like '2026 FIFA' if not data).
- If a number is illegible, omit it.
- layout.pass is true only if all five layout flags are true.
"""

AESTHETIC_SYSTEM = """You rate a public-safety/data infographic on visual quality ONLY.
Do NOT verify factual correctness.

Rate aesthetic_quality from 1 to 5:
5 = publication quality / professional agency
4 = good, minor issues
3 = acceptable, noticeable flaws
2 = poor, needs substantial rework
1 = unusable

Also rate readability 1-5 (legibility, contrast, label clarity).

Return ONLY JSON:
{"aesthetic_quality": <1-5>, "readability": <1-5>, "one_line_reason": "short"}
"""


def score_extract(client: OpenAI, path: Path) -> dict:
    parts = [
        {"type": "text", "text": "Extract numbers, charts, layout validity, and story elements from this infographic."},
        encode_image(path),
    ]
    data = vlm_json(client, EXTRACT_SYSTEM, parts, max_tokens=2500)
    # enforce layout pass
    layout = data.get("layout") or {}
    flags = ["full_render", "no_text_overflow", "no_illegal_overlap", "no_clipping", "structure_legible"]
    if all(k in layout for k in flags):
        layout["pass"] = all(bool(layout[k]) for k in flags)
    data["layout"] = layout
    story = data.get("story_elements") or {}
    keys = ["location", "time_or_period", "quantitative_evidence", "guidance_or_actions", "source"]
    count = sum(1 for k in keys if story.get(k))
    data["story_completeness_count"] = count
    data["story_completeness"] = round(count / 5.0, 4)
    return data


def score_aesthetic_median(client: OpenAI, path: Path, passes: int = AESTHETIC_PASSES) -> dict:
    scores_a, scores_r, reasons = [], [], []
    for i in range(passes):
        parts = [
            {"type": "text", "text": f"Rate aesthetic quality and readability (pass {i+1}/{passes})."},
            encode_image(path),
        ]
        # slight non-zero temp for variance reduction via median still; protocol says 3 passes
        data = vlm_json(client, AESTHETIC_SYSTEM, parts, max_tokens=300)
        try:
            scores_a.append(int(data.get("aesthetic_quality")))
        except (TypeError, ValueError):
            pass
        try:
            scores_r.append(int(data.get("readability")))
        except (TypeError, ValueError):
            pass
        reasons.append(data.get("one_line_reason", ""))
    return {
        "aesthetic_quality": int(statistics.median(scores_a)) if scores_a else None,
        "readability": int(statistics.median(scores_r)) if scores_r else None,
        "aesthetic_raw": scores_a,
        "readability_raw": scores_r,
        "reasons": reasons,
    }


# ---------------------------------------------------------------------------
# IO / resume
# ---------------------------------------------------------------------------

def load_done() -> set[tuple[str, str]]:
    done = set()
    if not JSONL.exists():
        return done
    for line in JSONL.read_text().splitlines():
        if not line.strip():
            continue
        try:
            r = json.loads(line)
            if r.get("ok"):
                done.add((r["case_id"], r["system"]))
        except json.JSONDecodeError:
            pass
    return done


def append_jsonl(rec: dict) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    with JSONL.open("a") as f:
        f.write(json.dumps(rec, ensure_ascii=False) + "\n")


def load_all() -> list[dict]:
    by = {}
    if not JSONL.exists():
        return []
    for line in JSONL.read_text().splitlines():
        if not line.strip():
            continue
        try:
            r = json.loads(line)
        except json.JSONDecodeError:
            continue
        if r.get("ok"):
            by[(r["case_id"], r["system"])] = r
    return [by[k] for k in sorted(by)]


def score_one(client: OpenAI, case_id: str, system: str) -> dict:
    path = image_path(case_id, system)
    if not path.exists():
        raise FileNotFoundError(path)
    reg = load_registry(case_id)
    gold = gold_numeric_values(reg)

    print(f"  extract...", flush=True)
    ext = score_extract(client, path)
    fid = score_fidelity(ext.get("visible_numbers") or [], gold)
    charts = ext.get("charts_detected") or []
    cr = chart_rule_pass(reg, charts)

    print(f"  aesthetics x{AESTHETIC_PASSES}...", flush=True)
    aes = score_aesthetic_median(client, path)

    layout_pass = bool((ext.get("layout") or {}).get("pass"))
    return {
        "ok": True,
        "case_id": case_id,
        "system": system,
        "image": str(path),
        "model_judge": MODEL,
        "ffr": fid["ffr"],
        "hr": fid["hr"],
        "numeric_accuracy": fid["numeric_accuracy"],
        "n_visible": fid["n_visible"],
        "n_matched": fid["n_matched"],
        "n_hallucinated": fid["n_hallucinated"],
        "fidelity_detail": fid,
        "chart_rule_pass": cr,
        "charts_detected": charts,
        "layout_pass": layout_pass,
        "layout_detail": ext.get("layout"),
        "story_completeness": ext.get("story_completeness"),
        "story_completeness_count": ext.get("story_completeness_count"),
        "story_elements": ext.get("story_elements"),
        "aesthetic_quality": aes.get("aesthetic_quality"),
        "readability": aes.get("readability"),
        "aesthetic_detail": aes,
        "extract_summary": ext.get("one_line_summary"),
    }


def aggregate(records: list[dict]) -> dict:
    by_sys: dict[str, list[dict]] = {s: [] for s in SYSTEMS}
    for r in records:
        by_sys.setdefault(r["system"], []).append(r)

    systems_out = {}
    for sys, rows in by_sys.items():
        if not rows:
            continue
        n = len(rows)
        def mean(key):
            vals = [r[key] for r in rows if r.get(key) is not None]
            return round(sum(vals) / len(vals), 4) if vals else None

        systems_out[sys] = {
            "n": n,
            "ffr": mean("ffr"),
            "hr": mean("hr"),
            "numeric_accuracy": mean("numeric_accuracy"),
            "chart_rule_pass_rate": round(
                100.0 * sum(r.get("chart_rule_pass", 0) for r in rows) / n, 1
            ),
            "layout_validity_pass_rate_pct": round(
                100.0 * sum(1 for r in rows if r.get("layout_pass")) / n, 1
            ),
            "story_completeness_mean_0_5": round(
                sum(r.get("story_completeness_count") or 0 for r in rows) / n, 2
            ),
            "aesthetic_quality_mean": mean("aesthetic_quality"),
            "readability_mean": mean("readability"),
            # for scatter: list of (ffr, aesthetic) per case
            "fidelity_x_aesthetics_points": [
                {"case_id": r["case_id"], "ffr": r["ffr"], "aesthetic": r["aesthetic_quality"]}
                for r in rows
            ],
        }

    return {
        "n_cases": len(CASES),
        "n_records": len(records),
        "judge_model": MODEL,
        "aesthetic_passes": AESTHETIC_PASSES,
        "skipped_metrics": [
            "Expert appropriateness (manual)",
            "Framework match distance (not in registry)",
            "Ablation no fact-locking (no outputs)",
        ],
        "notes": [
            "Story completeness adapted for data briefs: location, time/period, quantitative evidence, guidance, source (0-5).",
            "Chart-rule pass uses proxy data-shape heuristics, not official Phase-2 lookup.",
            "FFR/HR/NumericAccuracy: VLM transcription then deterministic gold match with tolerance.",
            "RiskScribe scored on infographic_decoration.png only.",
            f"Partial set: N={len(CASES)} of planned 20 complex cases.",
        ],
        "systems": systems_out,
    }


def write_csv(records: list[dict]) -> None:
    fields = [
        "case_id", "system", "ffr", "hr", "numeric_accuracy",
        "n_visible", "n_matched", "n_hallucinated",
        "chart_rule_pass", "layout_pass",
        "story_completeness_count", "story_completeness",
        "aesthetic_quality", "readability", "charts_detected",
    ]
    with CSV_PATH.open("w", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        for r in records:
            row = {k: r.get(k) for k in fields}
            row["charts_detected"] = "|".join(r.get("charts_detected") or [])
            w.writerow(row)


def fill_docx(summary: dict) -> None:
    from docx import Document

    if not DOCX_PATH.exists():
        print("DOCX missing; skip fill")
        return

    doc = Document(str(DOCX_PATH))
    if not doc.tables:
        print("No tables in docx")
        return
    t = doc.tables[0]  # Table 1

    # Map header cells to system keys
    header = [c.text.strip() for c in t.rows[0].cells]
    # Expect: Claim, Metric, Scale, Judge, GPT-Img-1.5, GPT-Img-2, CI sol, CI luna, RiskScribe
    col_for = {}
    for idx, h in enumerate(header):
        hl = h.lower()
        if "1.5" in hl or "img-1" in hl:
            col_for["gpt_img_1_5"] = idx
        elif "img-2" in hl or "image-2" in hl or "gpt-img-2" in hl:
            col_for["gpt_img_2"] = idx
        elif "sol" in hl:
            col_for["ci_sol"] = idx
        elif "luna" in hl:
            col_for["ci_luna"] = idx
        elif "riskscribe" in hl or "risk scribe" in hl:
            col_for["riskscribe"] = idx

    # Fallback fixed positions if headers odd
    if len(col_for) < 5 and len(header) >= 9:
        col_for = {
            "gpt_img_1_5": 4,
            "gpt_img_2": 5,
            "ci_sol": 6,
            "ci_luna": 7,
            "riskscribe": 8,
        }

    sys_data = summary.get("systems") or {}

    def cell(sys, text):
        if sys not in col_for:
            return
        # find row later

    # Row identification by metric text in col 1 (Metric)
    metric_row = {}
    for ri, row in enumerate(t.rows):
        if ri == 0:
            continue
        metric = row.cells[1].text.strip().lower() if len(row.cells) > 1 else ""
        claim = row.cells[0].text.strip().lower() if row.cells else ""
        key = metric + " " + claim
        if "fact-fidelity" in metric or metric.startswith("ffr") or "fact fidelity" in metric:
            metric_row["ffr"] = ri
        elif "hallucination" in metric:
            metric_row["hr"] = ri
        elif "numeric accuracy" in metric:
            metric_row["numeric_accuracy"] = ri
        elif "chart-rule" in metric or "chart rule" in metric:
            metric_row["chart_rule"] = ri
        elif "expert" in metric:
            metric_row["expert"] = ri
        elif "framework" in metric:
            metric_row["framework"] = ri
        elif "layout validity" in metric:
            metric_row["layout"] = ri
        elif "story completeness" in metric:
            metric_row["story"] = ri
        elif "aesthetic quality" in metric:
            metric_row["aesthetic"] = ri
        elif "fidelity" in metric and "aesthetic" in metric:
            metric_row["fxa"] = ri
        elif "ablation" in metric:
            metric_row["ablation"] = ri

    def set_row(row_key: str, formatter):
        ri = metric_row.get(row_key)
        if ri is None:
            return
        row = t.rows[ri]
        for sys, col in col_for.items():
            s = sys_data.get(sys)
            if not s:
                continue
            val = formatter(sys, s)
            if col < len(row.cells):
                row.cells[col].text = val

    set_row("ffr", lambda sys, s: f"{s['ffr']:.3f}" if s.get("ffr") is not None else "—")
    set_row("hr", lambda sys, s: f"{s['hr']:.3f}" if s.get("hr") is not None else "—")
    set_row(
        "numeric_accuracy",
        lambda sys, s: f"{s['numeric_accuracy']:.3f}" if s.get("numeric_accuracy") is not None else "—",
    )
    set_row("chart_rule", lambda sys, s: f"{s['chart_rule_pass_rate']}%")
    set_row("expert", lambda sys, s: "—")  # skipped
    set_row(
        "framework",
        lambda sys, s: "—" if sys == "riskscribe" else "—",
    )
    set_row("layout", lambda sys, s: f"{s['layout_validity_pass_rate_pct']}%")
    set_row("story", lambda sys, s: f"{s['story_completeness_mean_0_5']:.2f}")
    set_row(
        "aesthetic",
        lambda sys, s: f"{s['aesthetic_quality_mean']:.2f}"
        if s.get("aesthetic_quality_mean") is not None
        else "—",
    )

    def fxa_fmt(sys, s):
        pts = s.get("fidelity_x_aesthetics_points") or []
        # one summary point: mean FFR, mean aesthetic
        if s.get("ffr") is not None and s.get("aesthetic_quality_mean") is not None:
            return f"(FFR={s['ffr']:.2f}, Aes={s['aesthetic_quality_mean']:.1f})"
        return "—"

    set_row("fxa", fxa_fmt)
    set_row("ablation", lambda sys, s: "—" if sys == "riskscribe" else "—")

    out_copy = OUT / "RiskScribe_Final_Score_Tables_table1_filled.docx"
    doc.save(str(DOCX_PATH))
    doc.save(str(out_copy))
    print(f"Filled Table 1 in {DOCX_PATH}")
    print(f"Copy: {out_copy}")


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    client = OpenAI(api_key=load_key())
    done = load_done()
    print(f"Judge model: {MODEL}", flush=True)
    print(f"Already done: {len(done)}", flush=True)

    failures = []
    total = len(CASES) * len(SYSTEMS)
    i = 0
    for case_id in CASES:
        for system in SYSTEMS:
            i += 1
            if (case_id, system) in done:
                print(f"[{i}/{total}] SKIP {case_id}/{system}", flush=True)
                continue
            print(f"[{i}/{total}] SCORE {case_id}/{system}", flush=True)
            try:
                rec = score_one(client, case_id, system)
                append_jsonl(rec)
                print(
                    f"  -> FFR={rec['ffr']} HR={rec['hr']} chart={rec['chart_rule_pass']} "
                    f"layout={rec['layout_pass']} story={rec['story_completeness_count']}/5 "
                    f"aes={rec['aesthetic_quality']}",
                    flush=True,
                )
            except Exception as e:
                print(f"  FAILED: {e}", flush=True)
                append_jsonl({"ok": False, "case_id": case_id, "system": system, "error": str(e)})
                failures.append((case_id, system, str(e)))
                time.sleep(2)

    records = load_all()
    print(f"\nSuccessful records: {len(records)}/{total}", flush=True)
    if not records:
        return 1

    summary = aggregate(records)
    SUMMARY.write_text(json.dumps(summary, indent=2) + "\n")
    write_csv(records)
    fill_docx(summary)

    print("\n========== TABLE 1 AGGREGATE ==========", flush=True)
    for sys, s in summary["systems"].items():
        print(
            f"{sys:12s} FFR={s['ffr']:.3f} HR={s['hr']:.3f} "
            f"chart={s['chart_rule_pass_rate']}% layout={s['layout_validity_pass_rate_pct']}% "
            f"story={s['story_completeness_mean_0_5']:.2f} aes={s['aesthetic_quality_mean']}",
            flush=True,
        )
    print(json.dumps(summary, indent=2)[:3000], flush=True)
    if failures:
        print("Failures:", failures, flush=True)
    return 0 if not failures else 2


if __name__ == "__main__":
    sys.exit(main())
